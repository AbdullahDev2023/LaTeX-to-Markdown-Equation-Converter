import streamlit as st
import re
import base64
from io import BytesIO
import pandas as pd
from PIL import Image
import io
import markdown
import pdfkit
from html2image import Html2Image
import os
import json
import uuid
import time
import docx
from docx.shared import Pt
from streamlit_ace import st_ace
import datetime

# Initialize session state if needed
if 'history' not in st.session_state:
    st.session_state.history = []
if 'last_autosave' not in st.session_state:
    st.session_state.last_autosave = time.time()
if 'session_id' not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
if 'theme' not in st.session_state:
    st.session_state.theme = "light"

def convert_latex_to_markdown(text):
    """
    Convert ChatGPT-style LaTeX delimiters to Markdown-compatible delimiters.
    - \( ... \) → $ ... $
    - \[ ... \] → $$ ... $$
    """
    # Convert display equations: \[ ... \] to $$ ... $$
    text = re.sub(r'\\\[(.*?)\\\]', r'$$\1$$', text, flags=re.DOTALL)
    
    # Convert inline equations: \( ... \) to $ ... $
    text = re.sub(r'\\\((.*?)\\\)', r'$\1$', text, flags=re.DOTALL)
    
    return text

def markdown_to_html(markdown_text):
    """Convert markdown text to HTML."""
    return markdown.markdown(markdown_text, extensions=['extra', 'codehilite'])

def html_to_pdf(html_content, output_path="output.pdf"):
    """Convert HTML to PDF."""
    try:
        pdfkit.from_string(html_content, output_path)
        return output_path
    except Exception as e:
        st.error(f"Error converting to PDF: {e}")
        return None

def html_to_image(html_content, output_path="output.jpg"):
    """Convert HTML to image."""
    try:
        hti = Html2Image()
        hti.screenshot(html_str=html_content, save_as=output_path)
        return output_path
    except Exception as e:
        st.error(f"Error converting to image: {e}")
        return None

def get_download_link(file_path, link_text, file_type):
    """Generate a download link for a file."""
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        b64 = base64.b64encode(data).decode()
        href = f'<a href="data:application/{file_type};base64,{b64}" download="{os.path.basename(file_path)}">{link_text}</a>'
        return href
    except Exception as e:
        st.error(f"Error creating download link: {e}")
        return None

def save_to_local_storage():
    """Create JavaScript to save current input/output to localStorage"""
    current_data = {
        "user_input": st.session_state.user_input,
        "raw_output": st.session_state.raw_output,
        "display_mode": st.session_state.display_mode,
        "theme": st.session_state.theme,
        "last_saved": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    
    js_code = f"""
    <script>
        const data = {json.dumps(current_data)};
        localStorage.setItem('latex_converter_autosave', JSON.stringify(data));
        console.log('Data autosaved at: ' + data.last_saved);
        
        // Update autosave indicator if it exists
        const autosaveIndicator = document.getElementById('autosave-indicator');
        if (autosaveIndicator) {{
            autosaveIndicator.textContent = 'Last autosaved: ' + data.last_saved;
        }}
    </script>
    """
    
    st.markdown(js_code, unsafe_allow_html=True)
    st.session_state.last_autosave = time.time()
    
def load_from_local_storage():
    """Create JavaScript to load saved data from localStorage"""
    js_code = """
    <script>
        const savedData = localStorage.getItem('latex_converter_autosave');
        if (savedData) {
            const data = JSON.parse(savedData);
            
            // Create a form to submit the data back to Streamlit
            const form = document.createElement('form');
            form.method = 'POST';
            form.style.display = 'none';
            
            // Create hidden input for the data
            const input = document.createElement('input');
            input.type = 'hidden';
            input.name = 'localStorage_data';
            input.value = savedData;
            form.appendChild(input);
            
            // Create hidden submit button
            const submit = document.createElement('input');
            submit.type = 'submit';
            form.appendChild(submit);
            
            document.body.appendChild(form);
            form.submit();
        }
    </script>
    """
    return js_code

def add_to_history(input_text, output_text):
    """Add current conversion to history"""
    if len(input_text) > 0 and len(output_text) > 0:
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Truncate long entries for display
        input_preview = input_text[:100] + "..." if len(input_text) > 100 else input_text
        
        entry = {
            "timestamp": timestamp,
            "input": input_text,
            "input_preview": input_preview,
            "output": output_text
        }
        
        # Add to history if it's different from the last entry
        if not st.session_state.history or st.session_state.history[-1]["input"] != input_text:
            st.session_state.history.append(entry)
            
            # Keep history at maximum 20 items
            if len(st.session_state.history) > 20:
                st.session_state.history.pop(0)

def export_to_latex(markdown_text):
    """Convert markdown equations back to LaTeX format"""
    # Convert display equations: $$ ... $$ to \[ ... \]
    text = re.sub(r'\$\$(.*?)\$\$', r'\\[\1\\]', markdown_text, flags=re.DOTALL)
    
    # Convert inline equations: $ ... $ to \( ... \)
    text = re.sub(r'(?<!\$)\$(?!\$)(.*?)(?<!\$)\$(?!\$)', r'\\(\1\\)', text, flags=re.DOTALL)
    
    return text

def export_to_docx(markdown_text, output_path="output.docx"):
    """Export markdown to Word document"""
    try:
        # Create a new document
        doc = docx.Document()
        
        # Add heading
        doc.add_heading('Converted Document', 0)
        
        # Extract equations for special handling
        equations = {}
        eq_count = 0
        
        # Replace display equations with placeholders
        def replace_display_eq(match):
            nonlocal eq_count
            eq_id = f"DISPLAY_EQ_{eq_count}"
            eq_count += 1
            equations[eq_id] = match.group(1)
            return eq_id
        
        # Replace inline equations with placeholders
        def replace_inline_eq(match):
            nonlocal eq_count
            eq_id = f"INLINE_EQ_{eq_count}"
            eq_count += 1
            equations[eq_id] = match.group(1)
            return eq_id
        
        # Replace equations with placeholders
        text_with_placeholders = re.sub(r'\$\$(.*?)\$\$', replace_display_eq, markdown_text, flags=re.DOTALL)
        text_with_placeholders = re.sub(r'(?<!\$)\$(?!\$)(.*?)(?<!\$)\$(?!\$)', replace_inline_eq, text_with_placeholders, flags=re.DOTALL)
        
        # Split by paragraphs and add to document
        paragraphs = text_with_placeholders.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                # Check for equation placeholders
                for word in para.split():
                    if word.startswith("DISPLAY_EQ_") or word.startswith("INLINE_EQ_"):
                        p.add_run(f"[EQUATION: {equations[word]}]").italic = True
                    else:
                        p.add_run(word + " ")
        
        # Save the document
        doc.save(output_path)
        return output_path
    except Exception as e:
        st.error(f"Error exporting to Word: {e}")
        return None

def get_theme_styles():
    """Return CSS styles based on current theme"""
    light_theme = """
        .stApp {
            background-color: #f8f9fa;
            color: #212529;
        }
        .stTextArea textarea {
            background-color: #ffffff;
            font-family: 'Courier New', monospace;
        }
        .converted-text {
            background-color: #ffffff;
            padding: 20px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        h1, h2, h3 {
            color: #0066cc;
        }
        .info-box {
            background-color: #e6f3ff;
            padding: 15px;
            border-left: 5px solid #0066cc;
            border-radius: 5px;
            margin-bottom: 20px;
        }
        .export-button {
            background-color: #0066cc;
            color: white;
            border-radius: 5px;
            padding: 10px 15px;
            border: none;
            font-size: 14px;
            cursor: pointer;
            transition: all 0.3s;
            text-align: center;
            width: 100%;
            margin: 5px 0;
        }
        .export-button:hover {
            background-color: #004c99;
            transform: translateY(-2px);
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        }
        .history-item {
            border-left: 3px solid #0066cc; 
            padding: 12px;
            margin: 8px 0;
            background-color: #f9f9f9;
            cursor: pointer;
            border-radius: 5px;
            transition: all 0.2s;
        }
        .history-item:hover {
            background-color: #e6f3ff;
            transform: translateX(3px);
        }
    """
    
    dark_theme = """
        .stApp {
            background-color: #1a1a1a;
            color: #e0e0e0;
        }
        .stTextArea textarea {
            background-color: #2d2d2d;
            color: #e0e0e0;
            font-family: 'Courier New', monospace;
        }
        .converted-text {
            background-color: #2d2d2d;
            color: #e0e0e0;
            padding: 20px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.3);
        }
        h1, h2, h3 {
            color: #4d94ff;
        }
        .info-box {
            background-color: #2d333b;
            padding: 15px;
            border-left: 5px solid #4d94ff;
            border-radius: 5px;
            margin-bottom: 20px;
        }
        .export-button {
            background-color: #4d94ff;
            color: white;
            border-radius: 5px;
            padding: 10px 15px;
            border: none;
            font-size: 14px;
            cursor: pointer;
            transition: all 0.3s;
            text-align: center;
            width: 100%;
            margin: 5px 0;
        }
        .export-button:hover {
            background-color: #3a7bda;
            transform: translateY(-2px);
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.3);
        }
        .history-item {
            border-left: 3px solid #4d94ff; 
            padding: 12px;
            margin: 8px 0;
            background-color: #2d2d2d;
            cursor: pointer;
            border-radius: 5px;
            transition: all 0.2s;
        }
        .history-item:hover {
            background-color: #3a3a3a;
            transform: translateX(3px);
        }
    """
    
    common_styles = """
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stDeployButton {display:none;}
        .css-18e3th9 {
            padding-top: 1rem;
            padding-bottom: 1rem;
        }
        .css-1d391kg {
            padding: 1rem;
        }
        .title-container {
            text-align: center;
            margin-bottom: 2rem;
        }
        .offline-indicator {
            padding: 5px 10px;
            border-radius: 10px;
            font-size: 0.8em;
            display: inline-block;
            margin-top: 10px;
            animation: pulse 2s infinite;
        }
        @keyframes pulse {
            0% { opacity: 1; }
            50% { opacity: 0.7; }
            100% { opacity: 1; }
        }
        .autosave-indicator {
            font-size: 0.8em;
            text-align: right;
            font-style: italic;
            margin-top: 5px;
        }
        .ace-editor {
            border-radius: 10px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            margin-bottom: 20px;
        }
        .app-header {
            display: flex;
            align-items: center;
            justify-content: center;
            margin-bottom: 1.5rem;
        }
        .app-logo {
            font-size: 2.5rem;
            margin-right: 15px;
        }
        .theme-toggle {
            position: fixed;
            top: 70px;
            right: 20px;
            z-index: 1000;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            display: flex;
            align-items: center;
            justify-content: center;
            cursor: pointer;
            box-shadow: 0 2px 10px rgba(0,0,0,0.2);
            transition: all 0.3s;
        }
        .theme-toggle:hover {
            transform: rotate(30deg);
        }
        .btn-group {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            margin-bottom: 15px;
        }
        .btn-action {
            flex: 1;
            min-width: 100px;
            padding: 8px 16px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            transition: all 0.3s;
            font-weight: 500;
            text-align: center;
        }
        .btn-primary {
            background-color: #0066cc;
            color: white;
        }
        .btn-primary:hover {
            background-color: #004c99;
            transform: translateY(-2px);
        }
        .btn-secondary {
            background-color: #6c757d;
            color: white;
        }
        .btn-secondary:hover {
            background-color: #5a6268;
            transform: translateY(-2px);
        }
        .export-section {
            background-color: rgba(0,0,0,0.03);
            border-radius: 10px;
            padding: 15px;
            margin-top: 20px;
        }
        .export-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(150px, 1fr));
            gap: 10px;
            margin-top: 10px;
        }
        .section-title {
            font-size: 1.5rem;
            margin-bottom: 15px;
            padding-bottom: 8px;
            border-bottom: 2px solid rgba(0,102,204,0.3);
        }
        .card-container {
            background-color: rgba(255,255,255,0.05);
            border-radius: 10px;
            padding: 20px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            margin-bottom: 20px;
            transition: all 0.3s;
        }
        .card-container:hover {
            box-shadow: 0 6px 12px rgba(0,0,0,0.15);
            transform: translateY(-3px);
        }
        .tooltip {
            position: relative;
            display: inline-block;
            cursor: help;
        }
        .tooltip .tooltiptext {
            visibility: hidden;
            width: 200px;
            background-color: #555;
            color: #fff;
            text-align: center;
            border-radius: 6px;
            padding: 5px;
            position: absolute;
            z-index: 1;
            bottom: 125%;
            left: 50%;
            margin-left: -100px;
            opacity: 0;
            transition: opacity 0.3s;
        }
        .tooltip:hover .tooltiptext {
            visibility: visible;
            opacity: 1;
        }
    """
    
    return common_styles + (light_theme if st.session_state.theme == "light" else dark_theme)

def toggle_theme():
    """Toggle between light and dark theme"""
    if st.session_state.theme == "light":
        st.session_state.theme = "dark"
    else:
        st.session_state.theme = "light"
    
    # Save theme preference
    save_to_local_storage()

def main():
    # Set page config FIRST - before any other Streamlit commands
    st.set_page_config(
        page_title="LaTeX to Markdown Converter",
        page_icon="📝",
        layout="wide"
    )
    
    # Apply CSS styles based on current theme
    st.markdown(f'<style>{get_theme_styles()}</style>', unsafe_allow_html=True)
    
    # Theme toggle button
    theme_icon = "🌙" if st.session_state.theme == "light" else "☀️"
    st.markdown(
        f"""
        <div class="theme-toggle" onclick="window.location.href='?theme={('light' if st.session_state.theme == 'dark' else 'dark')}'">
            {theme_icon}
        </div>
        """, 
        unsafe_allow_html=True
    )
    
    # Check if theme was passed in URL
    params = st.query_params
    if "theme" in params and params["theme"][0] in ["light", "dark"]:
        if st.session_state.theme != params["theme"][0]:
            st.session_state.theme = params["theme"][0]
    
    # Title with styled container
    st.markdown('<div class="title-container">', unsafe_allow_html=True)
    st.markdown(
        """
        <div class="app-header">
            <div class="app-logo">📝</div>
            <h1>LaTeX to Markdown Equation Converter</h1>
        </div>
        """, 
        unsafe_allow_html=True
    )
    
    # Display offline indicator with styling based on theme
    offline_indicator_color = "#e8f5e9" if st.session_state.theme == "light" else "#2e3d2f"
    offline_indicator_text_color = "#2e7d32" if st.session_state.theme == "light" else "#4caf50"
    st.markdown(
        f'<div class="offline-indicator" style="background-color: {offline_indicator_color}; color: {offline_indicator_text_color};">✓ Works Offline</div>', 
        unsafe_allow_html=True
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    # App description
    st.markdown("""
    <div class="info-box">
    <strong>Convert LaTeX equations from ChatGPT-style formatting to Markdown-compatible formatting:</strong>
    <ul>
        <li><code>\\( ... \\)</code> → <code>$ ... $</code> (inline equations)</li>
        <li><code>\\[ ... \\]</code> → <code>$$ ... $$</code> (display equations)</li>
    </ul>
    <p>Perfect for copying equations from ChatGPT into Markdown documents or websites that use MathJax/KaTeX.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Create tabs for main app and history with custom styling
    tabs_css = """
    <style>
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px;
        }
        .stTabs [data-baseweb="tab"] {
            background-color: transparent;
            border-radius: 4px 4px 0 0;
            padding: 10px 16px;
            font-weight: 500;
        }
        .stTabs [aria-selected="true"] {
            background-color: rgba(0, 102, 204, 0.1);
            border-bottom: 2px solid #0066cc;
        }
    </style>
    """
    st.markdown(tabs_css, unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["✏️ Converter", "📋 History"])
    
    with tab1:
        # Example input section with improved styling
        with st.expander("📝 Show Example Input", expanded=False):
            example_text = """Here's an example of inline LaTeX: \\(E = mc^2\\) in a sentence.

And here's a display equation:
\\[
\\int_{a}^{b} f(x) \\, dx = F(b) - F(a)
\\]

You can have multiple equations in your text:
\\(\\alpha + \\beta = \\gamma\\) and \\(x^2 + y^2 = z^2\\)

Another display equation:
\\[
\\frac{d}{dx}\\left( \\int_{a}^{x} f(t) \\, dt \\right) = f(x)
\\]
"""
            st.code(example_text, language="markdown")
            if st.button("Use This Example", key="use_example"):
                st.session_state.user_input = example_text
                st.rerun()
        
        # User input section
        st.markdown('<div class="section-title">Input LaTeX</div>', unsafe_allow_html=True)
        
        # Initialize session state for user input if it doesn't exist
        if 'user_input' not in st.session_state:
            st.session_state.user_input = ""
        
        if 'raw_output' not in st.session_state:
            st.session_state.raw_output = ""
        
        if 'display_mode' not in st.session_state:
            st.session_state.display_mode = "side_by_side"  # Default to side-by-side
        
        # Add action buttons with improved styling
        st.markdown('<div class="btn-group">', unsafe_allow_html=True)
        col1, col2, col3 = st.columns([1, 1, 2])
        
        with col1:
            if st.button("📋 Paste from Clipboard", help="Paste content from your clipboard"):
                try:
                    import pyperclip
                    clipboard_text = pyperclip.paste()
                    st.session_state.user_input = clipboard_text
                    st.rerun()
                except ImportError:
                    st.error("Pyperclip not installed. Please install with: pip install pyperclip")
                except Exception as e:
                    st.error(f"Error accessing clipboard: {e}")
        
        with col2:
            if st.button("🔄 Load Autosave", help="Load your last automatically saved content"):
                st.markdown(load_from_local_storage(), unsafe_allow_html=True)
        
        with col3:
            st.markdown('<div id="autosave-indicator" class="autosave-indicator">Content will be autosaved</div>', unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Syntax highlighted editor with LaTeX support in a card container
        st.markdown('<div class="card-container">', unsafe_allow_html=True)
        st.markdown("### 📝 Input Text Editor")
        st.markdown('<div class="tooltip">Hover for tips <span class="tooltiptext">Paste LaTeX from ChatGPT with \\( \\) or \\[ \\] delimiters</span></div>', unsafe_allow_html=True)
        
        editor_theme = "github" if st.session_state.theme == "light" else "monokai"
        current_ace_value = st_ace(
            value=st.session_state.user_input,
            language="latex",
            theme=editor_theme,
            key="input_area",
            height=250,
            font_size=14,
            wrap=True,
            auto_update=True,
            show_gutter=True,
        )
        
        # Check if the value has changed
        if current_ace_value != st.session_state.user_input:
            st.session_state.user_input = current_ace_value
            if (time.time() - st.session_state.last_autosave) > 5:  # Autosave every 5 seconds
                save_to_local_storage()
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Convert LaTeX to Markdown in real-time
        if st.session_state.user_input:
            converted_text = convert_latex_to_markdown(st.session_state.user_input)
            st.session_state.raw_output = converted_text
            
            # Add to history
            add_to_history(st.session_state.user_input, converted_text)
            
            # Live Preview Section 
            st.markdown('<div class="section-title">Live Preview</div>', unsafe_allow_html=True)
            
            # Show rendered and raw output side by side
            col1, col2 = st.columns(2)
            
            # Rendered output column
            with col1:
                st.markdown('<div class="card-container">', unsafe_allow_html=True)
                st.markdown("### 🔍 Rendered Preview")
                st.markdown(f'<div class="converted-text">{converted_text}</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
                
            # Raw output column with editable text area
            with col2:
                st.markdown('<div class="card-container">', unsafe_allow_html=True)
                st.markdown("### 📄 Raw Output (Editable)")
                
                # Update callback for editable raw output
                def update_raw_output():
                    st.session_state.raw_output = st.session_state.editable_raw_output
                    if (time.time() - st.session_state.last_autosave) > 5:  # Autosave every 5 seconds
                        save_to_local_storage()
                    st.rerun()
                
                # Custom styling for text area
                text_area_height = 300
                text_area_bg_color = "#ffffff" if st.session_state.theme == "light" else "#2d2d2d"
                text_area_color = "#212529" if st.session_state.theme == "light" else "#e0e0e0"
                
                st.markdown(f"""
                <style>
                .stTextArea div[data-baseweb="textarea"] > div:first-child {{
                    height: {text_area_height}px;
                    background-color: {text_area_bg_color};
                    color: {text_area_color};
                    font-family: 'Courier New', monospace;
                    padding: 10px;
                    border-radius: 8px;
                }}
                </style>
                """, unsafe_allow_html=True)
                
                editable_raw = st.text_area(
                    "Edit converted markdown:",
                    value=st.session_state.raw_output,
                    key="editable_raw_output",
                    on_change=update_raw_output,
                    height=text_area_height
                )
                st.markdown('</div>', unsafe_allow_html=True)
            
            # Export options section with grid layout
            st.markdown('<div class="export-section">', unsafe_allow_html=True)
            st.markdown('<div class="section-title">Export Options</div>', unsafe_allow_html=True)
            
            export_formats = [
                {"name": "HTML", "icon": "🌐", "description": "Export as HTML with MathJax support"},
                {"name": "Markdown", "icon": "📝", "description": "Export as Markdown text file"},
                {"name": "PDF", "icon": "📄", "description": "Export as PDF document (requires wkhtmltopdf)"},
                {"name": "JPG", "icon": "🖼️", "description": "Export as JPG image (requires html2image)"},
                {"name": "LaTeX", "icon": "📐", "description": "Export back to LaTeX format"},
                {"name": "Word", "icon": "📘", "description": "Export as Word document"},
                {"name": "Plain Text", "icon": "📃", "description": "Export as plain text file"},
                {"name": "Copy to Clipboard", "icon": "📋", "description": "Copy converted text to clipboard"}
            ]
            
            st.markdown('<div class="export-grid">', unsafe_allow_html=True)
            
            # Generate export buttons
            for format_info in export_formats:
                export_name = format_info["name"]
                export_icon = format_info["icon"]
                export_desc = format_info["description"]
                
                if st.button(f"{export_icon} {export_name}", help=export_desc, key=f"export_{export_name.lower()}"):
                    if export_name == "HTML":
                        html_content = f"""
                        <!DOCTYPE html>
                        <html>
                        <head>
                            <title>Exported Markdown</title>
                            <meta charset="UTF-8">
                            <style>
                                body {{ 
                                    font-family: Arial, sans-serif; 
                                    padding: 20px; 
                                    max-width: 800px; 
                                    margin: 0 auto; 
                                    line-height: 1.6;
                                }}
                                pre {{ 
                                    background-color: #f5f5f5; 
                                    padding: 10px; 
                                    border-radius: 5px; 
                                    overflow-x: auto;
                                }}
                                code {{ font-family: 'Courier New', monospace; }}
                                .math-container {{
                                    padding: 10px 0;
                                    overflow-x: auto;
                                }}
                            </style>
                            <script src="https://cdnjs.cloudflare.com/ajax/libs/mathjax/2.7.7/MathJax.js?config=TeX-MML-AM_CHTML" async></script>
                        </head>
                        <body>
                            {markdown_to_html(st.session_state.raw_output)}
                        </body>
                        </html>
                        """
                        
                        b64 = base64.b64encode(html_content.encode()).decode()
                        href = f'<a href="data:text/html;base64,{b64}" download="converted_markdown.html" class="export-button">Download HTML</a>'
                        st.markdown(href, unsafe_allow_html=True)
                        st.success("HTML file ready for download! Click the button above to save it.")

                    elif export_name == "Markdown":
                        b64 = base64.b64encode(st.session_state.raw_output.encode()).decode()
                        href = f'<a href="data:text/markdown;base64,{b64}" download="converted_markdown.md" class="export-button">Download Markdown</a>'
                        st.markdown(href, unsafe_allow_html=True)
                        st.success("Markdown file ready for download! Click the button above to save it.")
                    
                    elif export_name == "PDF":
                        with st.spinner("Generating PDF..."):
                            try:
                                html_content = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <title>Exported Markdown</title>
                                    <meta charset="UTF-8">
                                    <style>
                                        body {{ 
                                            font-family: Arial, sans-serif; 
                                            padding: 20px; 
                                            max-width: 800px; 
                                            margin: 0 auto;
                                            line-height: 1.6;
                                        }}
                                        pre {{ 
                                            background-color: #f5f5f5; 
                                            padding: 10px; 
                                            border-radius: 5px;
                                            white-space: pre-wrap;
                                        }}
                                        code {{ font-family: 'Courier New', monospace; }}
                                    </style>
                                    <script src="https://cdnjs.cloudflare.com/ajax/libs/mathjax/2.7.7/MathJax.js?config=TeX-MML-AM_CHTML" async></script>
                                </head>
                                <body>
                                    {markdown_to_html(st.session_state.raw_output)}
                                </body>
                                </html>
                                """
                                
                                pdf_path = "converted_markdown.pdf"
                                pdfkit.from_string(html_content, pdf_path)
                                
                                with open(pdf_path, "rb") as pdf_file:
                                    pdf_bytes = pdf_file.read()
                                
                                b64 = base64.b64encode(pdf_bytes).decode()
                                href = f'<a href="data:application/pdf;base64,{b64}" download="converted_markdown.pdf" class="export-button">Download PDF</a>'
                                st.markdown(href, unsafe_allow_html=True)
                                st.success("PDF file ready for download! Click the button above to save it.")
                                
                                # Clean up
                                if os.path.exists(pdf_path):
                                    os.remove(pdf_path)
                                    
                            except Exception as e:
                                st.error(f"Error exporting to PDF: {e}")
                                st.info("Make sure wkhtmltopdf is installed on your system.")
                    
                    elif export_name == "JPG":
                        with st.spinner("Generating JPG image..."):
                            try:
                                html_content = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <title>Exported Markdown</title>
                                    <meta charset="UTF-8">
                                    <style>
                                        body {{ 
                                            font-family: Arial, sans-serif; 
                                            padding: 20px; 
                                            max-width: 800px; 
                                            margin: 0 auto;
                                            background-color: white;
                                            line-height: 1.6;
                                        }}
                                        pre {{ 
                                            background-color: #f5f5f5; 
                                            padding: 10px; 
                                            border-radius: 5px;
                                            white-space: pre-wrap;
                                        }}
                                        code {{ font-family: 'Courier New', monospace; }}
                                    </style>
                                    <script src="https://cdnjs.cloudflare.com/ajax/libs/mathjax/2.7.7/MathJax.js?config=TeX-MML-AM_CHTML" async></script>
                                </head>
                                <body>
                                    {markdown_to_html(st.session_state.raw_output)}
                                </body>
                                </html>
                                """
                                
                                # Save HTML to a temporary file
                                temp_html = "temp_export.html"
                                with open(temp_html, "w", encoding="utf-8") as f:
                                    f.write(html_content)
                                
                                # Convert to image
                                img_path = "converted_markdown.jpg"
                                hti = Html2Image()
                                hti.screenshot(url=os.path.abspath(temp_html), save_as=img_path)
                                
                                # Provide download link
                                with open(img_path, "rb") as img_file:
                                    img_bytes = img_file.read()
                                
                                b64 = base64.b64encode(img_bytes).decode()
                                href = f'<a href="data:image/jpeg;base64,{b64}" download="converted_markdown.jpg" class="export-button">Download JPG</a>'
                                st.markdown(href, unsafe_allow_html=True)
                                st.success("JPG image ready for download! Click the button above to save it.")
                                
                                # Clean up
                                if os.path.exists(temp_html):
                                    os.remove(temp_html)
                                if os.path.exists(img_path):
                                    os.remove(img_path)
                                    
                            except Exception as e:
                                st.error(f"Error exporting to JPG: {e}")
                                st.info("Make sure html2image is installed and a browser is available.")
                    
                    elif export_name == "LaTeX":
                        latex_content = export_to_latex(st.session_state.raw_output)
                        b64 = base64.b64encode(latex_content.encode()).decode()
                        href = f'<a href="data:text/plain;base64,{b64}" download="converted_latex.tex" class="export-button">Download LaTeX</a>'
                        st.markdown(href, unsafe_allow_html=True)
                        st.success("LaTeX file ready for download! Click the button above to save it.")
                    
                    elif export_name == "Word":
                        with st.spinner("Generating Word document..."):
                            try:
                                docx_path = "converted_markdown.docx"
                                export_to_docx(st.session_state.raw_output, docx_path)
                                
                                with open(docx_path, "rb") as docx_file:
                                    docx_bytes = docx_file.read()
                                
                                b64 = base64.b64encode(docx_bytes).decode()
                                href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="converted_markdown.docx" class="export-button">Download Word Document</a>'
                                st.markdown(href, unsafe_allow_html=True)
                                st.success("Word document ready for download! Click the button above to save it.")
                                
                                # Clean up
                                if os.path.exists(docx_path):
                                    os.remove(docx_path)
                                    
                            except Exception as e:
                                st.error(f"Error exporting to Word: {e}")
                                st.info("Make sure python-docx is installed with: pip install python-docx")
                    
                    elif export_name == "Plain Text":
                        # Remove markdown formatting for a plain text version
                        plain_text = st.session_state.raw_output
                        b64 = base64.b64encode(plain_text.encode()).decode()
                        href = f'<a href="data:text/plain;base64,{b64}" download="converted_plaintext.txt" class="export-button">Download Plain Text</a>'
                        st.markdown(href, unsafe_allow_html=True)
                        st.success("Plain text file ready for download! Click the button above to save it.")
                    
                    elif export_name == "Copy to Clipboard":
                        try:
                            import pyperclip
                            pyperclip.copy(st.session_state.raw_output)
                            st.success("Content copied to clipboard!")
                        except ImportError:
                            st.error("Pyperclip not installed. Please install with: pip install pyperclip")
                        except Exception as e:
                            st.error(f"Error copying to clipboard: {e}")
            
            st.markdown('</div>', unsafe_allow_html=True)  # Close export-grid
            
            # Additional export options or settings
            with st.expander("Advanced Export Settings", expanded=False):
                st.markdown("### Export Format Settings")
                
                st.markdown("#### PDF and Image Settings")
                col1, col2 = st.columns(2)
                with col1:
                    page_size = st.selectbox("Page Size", ["A4", "Letter", "Legal", "Tabloid"], index=0)
                with col2:
                    orientation = st.radio("Orientation", ["Portrait", "Landscape"], horizontal=True)
                
                st.markdown("#### Markdown Settings")
                preserve_newlines = st.checkbox("Preserve extra newlines", value=True)
                if not preserve_newlines and st.button("Compact Markdown", key="compact_md"):
                    # Remove excessive newlines
                    compacted = re.sub(r'\n{3,}', '\n\n', st.session_state.raw_output)
                    st.session_state.raw_output = compacted
                    st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)  # Close export-section
            
        else:
            st.info("Please enter some LaTeX text to convert. You can use the example above or paste from your clipboard.")
            
            # Show a quick demo if no input is provided
            with st.expander("Quick Demo", expanded=True):
                st.markdown("""
                ### How to Use This Tool
                
                1. Enter or paste text with LaTeX equations from ChatGPT or other sources
                2. See the live conversion to Markdown-compatible format
                3. Copy or export the result in various formats
                
                #### Example Conversion:
                
                **Original LaTeX:** `This is an equation: \\(E = mc^2\\)`
                
                **Converted:** `This is an equation: $E = mc^2$`
                
                Try the example button above to see more complex equations in action!
                """)
    
    # History tab
    with tab2:
        st.markdown('<div class="section-title">Conversion History</div>', unsafe_allow_html=True)
        
        if not st.session_state.history:
            st.info("No conversion history yet. Start converting LaTeX to see your history here.")
        else:
            st.write(f"Showing your last {len(st.session_state.history)} conversions:")
            
            # Clear history button with styling
            clear_btn_col1, clear_btn_col2 = st.columns([1, 5])
            with clear_btn_col1:
                if st.button("🗑️ Clear History", key="clear_history"):
                    st.session_state.history = []
                    st.rerun()
            
            # Search history
            with clear_btn_col2:
                search_term = st.text_input("🔍 Search history", placeholder="Type to filter history entries...")
            
            # Custom expander styling
            st.markdown("""
            <style>
            .streamlit-expanderHeader {
                background-color: rgba(0,102,204,0.05);
                border-radius: 5px;
            }
            </style>
            """, unsafe_allow_html=True)
            
            # Display history entries with filtering
            for i, entry in enumerate(reversed(st.session_state.history)):
                # Apply search filter if there's a search term
                if search_term and search_term.lower() not in entry['input'].lower() and search_term.lower() not in entry['output'].lower():
                    continue
                    
                with st.expander(f"🕒 {entry['timestamp']} - {entry['input_preview']}", expanded=False):
                    st.markdown("**Input:**")
                    st.code(entry['input'], language="latex")
                    
                    st.markdown("**Output:**")
                    st.code(entry['output'], language="markdown")
                    
                    # Action buttons for this history entry
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button(f"📋 Load This Entry", key=f"load_{i}"):
                            st.session_state.user_input = entry['input']
                            st.session_state.raw_output = entry['output']
                            st.rerun()
                    with col2:
                        if st.button(f"🗑️ Remove Entry", key=f"remove_{i}"):
                            reversed_index = len(st.session_state.history) - 1 - i
                            if 0 <= reversed_index < len(st.session_state.history):
                                st.session_state.history.pop(reversed_index)
                                st.rerun()
        
    # Footer with improved styling
    footer_bg_color = "#f8f9fa" if st.session_state.theme == "light" else "#1a1a1a"
    footer_text_color = "#666" if st.session_state.theme == "light" else "#999"
    
    st.markdown(f"""
    <div style="background-color: {footer_bg_color}; padding: 15px; margin-top: 30px; border-top: 1px solid rgba(0,0,0,0.1);">
        <div style="display: flex; justify-content: space-between; align-items: center; flex-wrap: wrap;">
            <div style="color: {footer_text_color}; font-size: 0.8em;">
                LaTeX to Markdown Converter | Made with ❤️ and Streamlit
            </div>
            <div style="color: {footer_text_color}; font-size: 0.8em;">
                <a href="#" style="color: {footer_text_color}; text-decoration: none; margin-right: 15px;">About</a>
                <a href="#" style="color: {footer_text_color}; text-decoration: none; margin-right: 15px;">Help</a>
                <a href="#" style="color: {footer_text_color}; text-decoration: none;">Feedback</a>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()